"""Upload router - 接入 parse_service 多模式解析，支持 image/pdf/txt/doc/excel"""
import uuid
import shutil
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from sqlalchemy.orm import Session

from backend.database.session import get_db
from backend.database.models import UploadRecord, IndicatorDefinition, IndicatorRecord
from backend.schemas.misc import UploadRecordOut
from backend.services.config_service import get_config

router = APIRouter(prefix="/api/upload", tags=["upload"])


def _get_uploads_dir() -> Path:
    cfg = get_config()
    d = Path(cfg.upload.dir)
    if not d.is_absolute():
        d = Path(__file__).resolve().parent.parent.parent / cfg.upload.dir
    d.mkdir(parents=True, exist_ok=True)
    return d


def _classify_file_type(ext: str) -> str:
    if ext in {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif"}:
        return "image"
    if ext == ".pdf":
        return "pdf"
    if ext in {".doc", ".docx"}:
        return "doc"
    if ext in {".xls", ".xlsx", ".csv"}:
        return "excel"
    return "text"


# ── 从 doc/docx 提取文字 ──────────────────────────────────────────────────────

def _extract_text_from_doc(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append("\t".join(cells))
        return "\n".join(lines)
    except ImportError:
        raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")


def _extract_text_from_excel(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".csv":
        import csv
        rows = []
        with open(file_path, "r", encoding="utf-8-sig", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append("\t".join(row))
        return "\n".join(rows)
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            rows.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append("\t".join(cells))
        return "\n".join(rows)
    except ImportError:
        # 降级到 xlrd（仅 .xls）
        try:
            import xlrd
            wb = xlrd.open_workbook(file_path)
            rows = []
            for sheet in wb.sheets():
                rows.append(f"[Sheet: {sheet.name}]")
                for rx in range(sheet.nrows):
                    cells = [str(sheet.cell_value(rx, cx)) for cx in range(sheet.ncols)]
                    rows.append("\t".join(cells))
            return "\n".join(rows)
        except ImportError:
            raise RuntimeError("openpyxl 未安装，请运行: pip install openpyxl")


# ── Routes ────────────────────────────────────────────────────────────────────

@router.post("/file", response_model=UploadRecordOut)
async def upload_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """上传文件并保存，返回记录 ID。支持 jpg/png/pdf/doc/docx/xls/xlsx/csv/txt"""
    ext = Path(file.filename or "unknown").suffix.lower()
    file_id = str(uuid.uuid4())
    saved_name = f"{file_id}{ext}"
    dest = _get_uploads_dir() / saved_name

    with dest.open("wb") as f:
        shutil.copyfileobj(file.file, f)

    file_type = _classify_file_type(ext)

    record = UploadRecord(
        id=file_id,
        file_path=str(dest),
        file_name=file.filename,
        file_type=file_type,
        status="pending",
    )
    db.add(record)
    db.commit()
    db.refresh(record)
    return record


@router.post("/analyze/{upload_id}", response_model=UploadRecordOut)
async def analyze_upload(upload_id: str, db: Session = Depends(get_db)):
    """
    触发 AI 解析：根据文件类型自动选择解析方式。
    image → OCR + LLM 或 vision 模型
    pdf   → pymupdf 提取文字 + LLM，失败则转图片
    doc/docx → python-docx 提取文字 + LLM
    excel/xls/xlsx/csv → openpyxl 提取表格 + LLM
    text  → 直接 LLM
    """
    record = db.query(UploadRecord).filter(UploadRecord.id == upload_id).first()
    if not record:
        raise HTTPException(404, "文件记录不存在")
    if not record.file_path or not Path(record.file_path).exists():
        raise HTTPException(400, "文件不存在，无法解析")

    from backend.services.parse_service import parse_lab_image, parse_lab_text, parse_lab_document

    record.status = "processing"
    db.commit()

    try:
        if record.file_type == "image":
            result = parse_lab_image(record.file_path)
        elif record.file_type == "pdf":
            result = parse_lab_document(record.file_path)
        elif record.file_type == "doc":
            text = _extract_text_from_doc(record.file_path)
            record.raw_ocr_text = text
            if not text.strip():
                result = {"indicators": [], "confidence": 0.0, "error": "文档无文字内容"}
            else:
                result = parse_lab_text(text)
        elif record.file_type == "excel":
            text = _extract_text_from_excel(record.file_path)
            record.raw_ocr_text = text
            if not text.strip():
                result = {"indicators": [], "confidence": 0.0, "error": "表格无内容"}
            else:
                result = parse_lab_text(text)
        else:  # text
            with open(record.file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            record.raw_ocr_text = content
            result = parse_lab_text(content)

        record.ai_parsed_json = result
        record.status = "done"
    except Exception as e:
        record.status = "failed"
        record.error_msg = str(e)

    db.commit()
    db.refresh(record)
    return record


@router.post("/confirm/{upload_id}")
async def confirm_upload(
    upload_id: str,
    items: Optional[List[dict]] = None,   # 前端可传入编辑后的指标列表（覆盖 ai_parsed_json）
    db: Session = Depends(get_db),
):
    """
    将解析结果批量写入 indicator_records。
    - items 参数：前端预览确认后可传回编辑过的指标列表；不传则使用 ai_parsed_json 原始结果
    - 匹配策略：优先按 indicator_id（前端已匹配）> code > 名称模糊
    """
    record = db.query(UploadRecord).filter(UploadRecord.id == upload_id).first()
    if not record:
        raise HTTPException(404, "文件记录不存在")
    if record.status != "done" or not record.ai_parsed_json:
        raise HTTPException(400, "尚未完成解析，无法确认入库")

    parsed = record.ai_parsed_json
    indicators_data = items if items is not None else parsed.get("indicators", [])
    report_date = parsed.get("report_date")

    all_defs = {d.code.lower(): d for d in db.query(IndicatorDefinition).all()}
    all_defs_by_name = {d.name: d for d in db.query(IndicatorDefinition).all()}

    imported = 0
    skipped = []
    auto_created = []

    for item in indicators_data:
        if not item.get("name") and not item.get("code"):
            continue

        # 优先用前端传回的 indicator_id（已手动匹配）
        defn = None
        if item.get("indicator_id"):
            defn = db.query(IndicatorDefinition).filter(
                IndicatorDefinition.id == item["indicator_id"]
            ).first()

        code = (item.get("code") or "").strip()
        name = (item.get("name") or "").strip()

        if not defn and code:
            defn = all_defs.get(code.lower())
        if not defn and name:
            # 模糊匹配名称
            for def_name, d in all_defs_by_name.items():
                if name in def_name or def_name in name:
                    defn = d
                    break

        # 仍未匹配：自动创建自定义指标（前端已勾选"自动创建"选项时）
        if not defn and item.get("auto_create"):
            from datetime import datetime
            defn = IndicatorDefinition(
                id=str(uuid.uuid4()),
                name=name or code,
                code=code or name,
                unit=item.get("unit"),
                category="其他",
                is_system=False,
                created_at=datetime.utcnow(),
            )
            db.add(defn)
            db.flush()
            auto_created.append(name or code)

        if not defn:
            skipped.append(name or code)
            continue

        rec_date = item.get("recorded_at") or report_date
        if not rec_date:
            skipped.append(f"{name}(无日期)")
            continue

        obj = IndicatorRecord(
            id=str(uuid.uuid4()),
            indicator_id=defn.id,
            value=item.get("value"),
            value_text=item.get("value_text"),
            recorded_at=rec_date,
            source_type="ai",
            source_ref=upload_id,
            note=f"自动导入自 {record.file_name}",
        )
        db.add(obj)
        imported += 1

    db.commit()
    return {
        "imported": imported,
        "skipped": skipped,
        "auto_created": auto_created,
        "message": (
            f"成功导入 {imported} 条"
            + (f"，新建指标 {len(auto_created)} 个" if auto_created else "")
            + (f"，跳过 {len(skipped)} 条（未匹配/无日期）" if skipped else "")
        ),
    }


@router.get("/records", response_model=list[UploadRecordOut])
def list_uploads(db: Session = Depends(get_db)):
    return db.query(UploadRecord).order_by(UploadRecord.created_at.desc()).all()


@router.delete("/{upload_id}")
def delete_upload(upload_id: str, db: Session = Depends(get_db)):
    obj = db.query(UploadRecord).filter(UploadRecord.id == upload_id).first()
    if not obj:
        raise HTTPException(404, "记录不存在")
    if obj.file_path and Path(obj.file_path).exists():
        Path(obj.file_path).unlink(missing_ok=True)
    db.delete(obj)
    db.commit()
    return {"ok": True}
